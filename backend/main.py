import faiss
import numpy as np
import io
from fastapi import FastAPI, File, UploadFile, Request, HTTPException 
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from starlette.responses import JSONResponse
from azure.storage.blob import BlobServiceClient
import os
from dotenv import load_dotenv
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import pdb
from datetime import datetime
import re
import requests

from BasePrompts import eti_prompt


from BlogExamples import blog_examples, recent_example, stamos_example, disney_example, long_form_examples, sector_example
import PyPDF2
# This is a minor change to trigger redeployment

RESULTS_DIR = os.getenv('RESULTS_DIR', 'results')


print('GOT EVERYTIHG LOADED SUCESFULLY')


load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")
account_key = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")

client = openai.OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
)

app = FastAPI()

# Configure CORS
origins = [
    "http://localhost:3000",  # React dev server
    "https://askyorkville-c3ckc8hgh4hzajeu.eastus-01.azurewebsites.net"  # For HTTPS
]

app.add_middleware(
    CORSMiddleware,
    #allow_origins=origins,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

print('GOT CORS CONFIGURED SUCCESSFULLY')


chunks_storage = {}  # {index: chunk_text}

#promptTemplates = ["Ad Hoc Query", "Tear Sheet", "Long Form", "SuperLong", "One Page Current Events"]



def get_chunk_by_index(idx):
    return chunks_storage.get(idx, "")

@app.exception_handler(413)
async def request_entity_too_large_handler(request: Request, exc: Exception):
    return JSONResponse(
        status_code=413,
        content={"message": "Payload too large"},
    )

# Increase the maximum payload size
from fastapi.middleware.trustedhost import TrustedHostMiddleware
app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=["*"],
)

# Initialize Faiss index
d = 1536  # Dimension of the embeddings
index = faiss.IndexFlatL2(d)

# Initialize Azure Blob Storage
account_name = "yorkvilleworks9016610742"

# Form the connection string
connection_string = f"DefaultEndpointsProtocol=https;AccountName={account_name};AccountKey={account_key};EndpointSuffix=core.windows.net"

# Create a BlobServiceClient
print("Creating BlobServiceClient...")
blob_service_client = BlobServiceClient.from_connection_string(connection_string)

container_name = "uploaded-files"

def split_text_into_chunks(text, max_tokens=8192):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        current_chunk.append(word)
        if len(current_chunk) >= max_tokens:
            chunks.append(' '.join(current_chunk))
            current_chunk = []

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def get_embeddings(text, chunk_size=8192):
    text_chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    
    embeddings = []
    for chunk in text_chunks:
        response = client.embeddings.create(
            input=chunk,
            model="text-embedding-3-small"
        )
        embedding = np.array(response.data[0].embedding).astype('float32')
        embeddings.append(embedding)
    
    embeddings = np.vstack(embeddings)
    
    return embeddings, text_chunks





def extract_text_from_pdf(file_location):
    with open(file_location, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def get_chunks_from_db(db_name):
    index_blob_client = blob_service_client.get_blob_client(container=container_name, blob=db_name)
    index_data = index_blob_client.download_blob().readall()
    temp_index_file = f"temp_{db_name}"
    
    with open(temp_index_file, "wb") as f:
        f.write(index_data)
    
    temp_index = faiss.read_index(temp_index_file)
    embeddings = temp_index.reconstruct_n(0, temp_index.ntotal)
    os.remove(temp_index_file)
    
    # Retrieve the chunks associated with this database
    chunks_blob_name = db_name.replace('_index', '_chunks')
    print(f"Fetching chunks for {chunks_blob_name}")
    
    chunks_blob_client = blob_service_client.get_blob_client(container=container_name, blob=chunks_blob_name)
    chunks_data = chunks_blob_client.download_blob().readall().decode('utf-8')
    chunks = chunks_data.split('\n')
    
    return chunks, embeddings




def generate_section(prompt, model_name):
    #response = openai.ChatCompletion.create(
    #    model=model_name,
    #    messages=[
    #        {"role": "system", "content": "You are a helpful assistant."},
    #        {"role": "user", "content": prompt}
    #    ],
    #    max_tokens=3500
    #)
    
    response =client.chat.completions.create(
        messages=[
            {"role": "system", "content": "You are a business, strategy, and finance expert, as well as a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        model = model_name,
        temperature = .3,
        max_tokens=3500
    )
    
    
    #return response.choices[0].message['content']
    return response.choices[0].message.content


def add_formatted_content(doc, formatted_text):
    lines = formatted_text.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            para = doc.add_paragraph()
            run = para.add_run(line[6:])
            run.bold = True
        elif line.startswith('- '):
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(line[2:])
        elif len(line) > 3 and line[0].isdigit() and line[1] == '.' and line[2] == ' ':
            para = doc.add_paragraph(style='List Number')
            para.add_run(line[3:])
        else:
            para = doc.add_paragraph()
            para.add_run(line)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def analyze_text_with_gpt(text):
    
    #prompt=f"Analyze the following text and determine which sections should be formatted as headings, subheadings, paragraphs, or lists. Return a structured representation of the text:\n\n{text}",
   
    
    
    jls_extract_var = [
            #{"role": "system", "content": f"You are a helpful assistant. in this case, you are helping a user create blog posts from summaries of text transcripts. The user will provide you with examples of original text summary and final blog post. then the user will provide an additional original text summary. your task is to create a blog post from this new original text summary.  it should be similar in style and approach as the given examples. note it is good if the blog posts have section headings to organize the thinking. please make the contents of each section quite a lot longer than current example sections are. The purpose of the blog series is to provide insights into happenings in financial markets and products used to trade and invest in these markets.Please don't use overly hyperbolic language, but definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone should be more cautiously optimisitc, where optimisim is appropriate, as befits analyses related to potential investments. Please don't use overly hyperbolic language, but definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone should be more cautiously optimisitc, where optimisim is appropriate, as befits analyses related to potential investments. also , here are a few other points to keep in mind: {eti_prompt}. Also please try not to use the word 'delve'.  "}
       {"role": "system", "content": f"""
       You are a helpful assistant and a meticulous editor
       """}
       
        ]
    #jls_extract_var.append({"role": "user", "content": f"here are some examples: {blog_examples}"})
    jls_extract_var.append({"role": "user", "content": f""""
    Analyze the following text and determine which sections should be formatted as headings, subheadings, paragraphs, or lists.
     In no event should the following symbols and text appear : "```markdown" or the word "markdown" which I have seen in some past examples.  also I don't want any asterisks '*' within bullet points or numbered lists. in the text nor any instances of the word 'bold' or 'italic'.
     also please no instances of  " ``` " which I have seen in some past examples.
     I do not like to have more than one line of blank space between any two sections.
    In addition, on some occasions I have seen numbered lists where the numbering of each item appears twice.  It is as though in some cases the original text has a numbered list and then this formatting exercise decides this text is a numbered list and so numbers it again.  please avoid this.
     
     Please return a structured representation of the text:\n\n{text}"""})  
        



    response = client.chat.completions.create(
        messages=jls_extract_var,
        model = "gpt-4o",
                        temperature = .3,
                        
                    
                        max_tokens=2500
    )



    
    response_str = response.choices[0].message.content

    

    return response_str
    



def clean_document(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = clean_text(run.text)



def clean_text(text):
    # Remove unwanted symbols like '*', '#', and other non-alphanumeric characters
    cleaned_text = re.sub(r'[^\w\s,.()-]', '', text)  # Keep alphanumeric, spaces, and some punctuation
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Remove extra spaces
    return cleaned_text


def sequential_superlong(documents, company_name, model_name):
   

    today_date = datetime.today().strftime('%Y-%m-%d_%H-%M-%S')

    #sanitized_company_name = re.sub(r'\s+', '_', company_name).replace("'", "").replace('"', "")
    sanitized_company_name = re.sub(r'\s+', '_', company_name).replace("'", "").replace('"', "")
    #filename = f"Formatted_SuperLong_{today_date}_{sanitized_company_name}.docx"
    filename = f"Formatted_SuperLong_{today_date}.docx"
    print("sanitized_company_name:",sanitized_company_name)
    #filename = f"Formatted_SuperLong_{today_date}_{company_name}.docx"
    #docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"
    docx_file_path = os.path.join(RESULTS_DIR, filename)

 

# Ensure no unwanted quotes
    docx_file_path = docx_file_path.replace('"', '').replace("'", "")


    ## Create a new Document
    #doc = Document()

    ## Add some dummy content
    #doc.add_heading('Dummy Document', 0)
    #doc.add_paragraph(f'This is a dummy document for {company_name}.')
    #doc.add_paragraph('This is just for testing the file writing and downloading process.')

    ## Save the document
    ##doc.save(docx_file_path)

    ##print(f"Dummy document saved to {docx_file_path}")

    #return docx_file_path
    
 
    
    all_str = " "


    text_so_far = " "

    today_date = datetime.today().strftime('%Y-%m-%d')

# Section Prompts
    prompts = {
        

        "company_analysis": f"""
            You are an expert business analyst.You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company Based on the following documents, generate a 3-page Company Analysis for {company_name}.
            Provide an overview of the company’s history and business model.
            
            
            
            
            Start with an overview of the company's Origin, History , Development and Current Business Model
            
            Analyze the financial performance including key metrics, growth, and profitability.
            Please make sure to capture and cite any information related to activist shareholders or investors being involved.  This is always interesting.  Also please make sure to mention any stock buyback or repurchase activity the company may be involved in.
            Please also Perform a SWOT analysis (Strengths, Weaknesses, Opportunities, Threats). Again, in this section, please only include information from the documents you are being presented with.  Do not search the web or use other sources for this section or the SWOT analysis
            
            Please write in an engaging and informative style, highlighting the most important points but don't make it too terse. 
            We do want a lot of data and statistics drawn for the text,  but we also want some narrative.
            This should be an informative and accurate read but something we'd be able to listent to as a podcast episode and not get bored.
            Do the best you can to execute the analyses according to the framework outlined.  but if the supplied text won't allow it. then gracefuly avoid the requested sections and move on.

            
            Please try to make sure your information is as accurate and up to date as possible at least as up to date as two months ago. Note today's date is {today_date} If you are not sure about something, please don't include it.  But if you are sure, please include it.
            Explicitly avoid talking about divisions that may have recently been sold or operations that have recent ceased. For example, Johnson and Johnson sold off its consumer health division in August 2023.  So don’t talk about this business or it’s consumers in the analysis 
            
            Also, please don't present in just bullet points.  while we like lots of specific data, we also like prose.  an example of the style we are going after is here: {disney_example}
            but not this example is just for style reference. But I do not want any overall summarizing or concluding paragraph in this section, notwithstanding whether the example has this. do not include any of the content from this example in the summary you are creating here.

                        
            For this section, please make sure to only include information about {company_name} that comes form the documents you are being presented as part of this prompt.  Do not search the web or use other sources for this section. 
            You may search the web for information about the industry and competitors as you create these sections.
            That said, while I don't want facts or data about {company_name} in this section coming from outside the documents, it is ok to use your strong business acumen and reasoning ability when creating this section.    
            When organizing the seciton, please feel free to use sub-numbering where it makes sense.  This is section 2 of the overall document, so if it makes sense you could talk baout some potins as 2.1, 2.2, etc.  But please don't force things.
                        
            
            Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
           
            Please make sure not repeat anything you've already produced for this document, which is: {text_so_far}

            Also, please don't include any repetitive or similar section headings at the begining of each section. I've seen a lot of this in previous versions.  
            One section title is fine.  then distinct subheadings, if necessary through the text.  but please avoid repetition.

            Please try not to use the word "delve".

            Here are the source Documents:
            {documents}
            """,

        "customer_analysis": f"""
            
            
            
            You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. Based on the following documents, generate a 2-3 page Customer Analysis for {company_name}.
            Analyze customer demographics and segmentation.
            Describe the needs and preferences of these customers.
            Evaluate customer satisfaction and loyalty.

            Please write in an engaging and informative style, highlighting the most important points but don't make it too terse. 
            We do want a lot of data and statistics drawn for the text,  but we also want some narrative.
            This should be an informative and accurate read but something we'd be able to listent to as a podcast episode and not get bored.
            Do the best you can to execute the analyses according to the framework outlined.  but if the supplied text won't allow it. then gracefuly avoid the requested sections and move on.

           
            
            Also, please don't present in just bullet points.  while we like lots of specific data, we also like prose.  an example of the style we are going after is here: {disney_example}
            but not this example is just for style reference. But I do not want any overall summarizing or concluding paragraph in this section, notwithstanding whether the example has this.  do not include any of the content from this example in the summary you are creating here.

            Please try to make sure your information is as accurate and up to date as possible at least as up to date as two months ago. Note today's date is {today_date} If you are not sure about something, please don't include it.  But if you are sure, please include it.

            Explicitly avoid talking about divisions that may have recently been sold or operations that have recent ceased. For example, Johnson and Johnson sold off its consumer health division in August 2023.  So don’t talk about this business or it’s consumers in the analysis 
                          
            For this section, please make sure to only include information about {company_name} that comes form the documents you are being presented as part of this prompt.  Do not search the web or use other sources for this section. 
            You may search the web for information about the industry and competitors as you create these sections.
            That said, while I don't want facts or data about {company_name} in this section coming from outside the documents, it is ok to use your strong business acumen and reasoning ability when creating this section.    
            When organizing the seciton, please feel free to use sub-numbering where it makes sense.  This is section 2 of the overall document, so if it makes sense you could talk baout some potins as 2.1, 2.2, etc.  But please don't force things.
                     




            Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.

            Please make sure not repeat anything you've already produced for this document, which is: {text_so_far}
            
            Also, please don't include any repetitive or similar section headings at the begining of each section. I've seen a lot of this in previous versions.  
            One section title is fine.  then distinct subheadings, if necessary through the text.  but please avoid repetition.

             Please try not to use the word "delve".



            Here are the source Documents:
            {documents}
            """,

        "competitor_analysis": f"""
            You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. Based on the following documents, generate a 2-3 page Competitor Analysis for {company_name}.
            Identify and analyze the major competitors and their market share.
            Describe their competitive strategies and positioning.
            Provide a comparative performance analysis.

            Also, please mention developments in the sector that might impact the competitive landscape.

            For this section when discussing {company_name}, please make sure to only include information about {company_name} that comes form the documents you are being presented as part of this prompt.  Do not search the web or use other sources regarding {company_name} for this section. 
            For information on other competitor ompanies you may use sources outside the documents share as part of this request.  Still, please only present information you are quite sure is factual. Do not make things up.


            Please write in an engaging and informative style, highlighting the most important points but don't make it too terse. 
            We do want a lot of data and statistics drawn for the text,  but we also want some narrative.
            This should be an informative and accurate read but something we'd be able to listent to as a podcast episode and not get bored.
            Do the best you can to execute the analyses according to the framework outlined.  but if the supplied text won't allow it. then gracefuly avoid the requested sections and move on.

            
            Please try to make sure your information is as accurate and up to date as possible at least as up to date as two months ago. Note today's date is {today_date} If you are not sure about something, please don't include it.  But if you are sure, please include it.
            
            Explicitly avoid talking about divisions that may have recently been sold or operations that have recent ceased. For example, Johnson and Johnson sold off its consumer health division in August 2023.  So don’t talk about this business or it’s consumers in the analysis 
                          
            For this section, please make sure to only include information about {company_name} that comes form the documents you are being presented as part of this prompt.  Do not search the web or use other sources for this section. 
            You may search the web for information about the industry and competitors as you create these sections.
            That said, while I don't want facts or data about {company_name} in this section coming from outside the documents, it is ok to use your strong business acumen and reasoning ability when creating this section.    
            When organizing the seciton, please feel free to use sub-numbering where it makes sense.  This is section 2 of the overall document, so if it makes sense you could talk baout some potins as 2.1, 2.2, etc.  But please don't force things.
                     


            Also, please don't present in just bullet points.  while we like lots of specific data, we also like prose.  an example of the style we are going after is here: {disney_example}
            but not this example is just for style reference. But I do not want any overall summarizing or concluding paragraph in this section, notwithstanding whether the example has this. do not include any of the content from this example in the summary you are creating here.

            
            
            
           Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
           

            Please make sure not repeat anything you've already produced for this document, which is: {text_so_far}

            Also, please don't include any repetitive or similar section headings at the begining of each section. I've seen a lot of this in previous versions.  
            One section title is fine.  then distinct subheadings, if necessary through the text.  but please avoid repetition.

             Please try not to use the word "delve".

            
            Here are the source Documents:
            {documents}
            """,

        "porters_five_forces": f"""
            You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. Based on the following documents, generate a 5-page Porter’s Five Forces Analysis for {company_name}.
            - Competitive Rivalry: Discuss the market structure and competitive landscape.
            - Supplier Power: Identify key suppliers, their bargaining power, and supplier relationships.
            - Buyer Power: Identify key buyers, their bargaining power, and buyer behavior.
            - Threat of Substitution: Identify alternative products/services and evaluate substitution risk.
            - Threat of New Entry: Discuss barriers to entry and potential new entrants.

            Please write in an engaging and informative style, highlighting the most important points but don't make it too terse. 
            We do want a lot of data and statistics drawn for the text,  but we also want some narrative.
            This should be an informative and accurate read but something we'd be able to listent to as a podcast episode and not get bored.
            Do the best you can to execute the analyses according to the framework outlined.  but if the supplied text won't allow it. then gracefuly avoid the requested sections and move on.

            Also, please don't present in just bullet points.  while we like lots of specific data, we also like prose.  an example of the style we are going after is here: {disney_example}
            but not this example is just for style reference.  But I do not want any overall summarizing or concluding paragraph in this section, notwithstanding whether the example has this. do not include any of the content from this example in the summary you are creating here.

            Please try to make sure your information is as accurate and up to date as possible at least as up to date as two months ago. Note today's date is {today_date} If you are not sure about something, please don't include it.  But if you are sure, please include it.

            Explicitly avoid talking about divisions that may have recently been sold or operations that have recent ceased. For example, Johnson and Johnson sold off its consumer health division in August 2023.  So don’t talk about this business or it’s consumers in the analysis 
                          
            For this section, please make sure to only include information about {company_name} that comes form the documents you are being presented as part of this prompt.  Do not search the web or use other sources for this section. 
            You may search the web for information about the industry and competitors as you create these sections.
            That said, while I don't want facts or data about {company_name} in this section coming from outside the documents, it is ok to use your strong business acumen and reasoning ability when creating this section.    
            When organizing the seciton, please feel free to use sub-numbering where it makes sense.  This is section 2 of the overall document, so if it makes sense you could talk baout some potins as 2.1, 2.2, etc.  But please don't force things.
                     


            Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.

            Please make sure not repeat anything you've already produced for this document, which is: {text_so_far}
            
            Also, please don't include any repetitive or similar section headings at the begining of each section. I've seen a lot of this in previous versions.  
            One section title is fine.  then distinct subheadings, if necessary through the text.  but please avoid repetition.

             Please try not to use the word "delve".


            Here are the source Documents:
            {documents}
            """,

    }

   
# Create a new Document
    doc = Document()
 
    # Generate each section and store the results
    report_sections = {}
    
    print('Inside sequential_superlong. about to loop through prompts')
    for section, prompt in prompts.items():
        report_sections[section] = generate_section(prompt, model_name)
        report_sections[section] = clean_text(report_sections[section])
        formatted_section = analyze_text_with_gpt(report_sections[section])    
        add_formatted_content(doc, formatted_section)
        text_so_far = text_so_far + '\n' + report_sections[section]
        print('Inside sequential_superlong. did a loop')
        #pdb.set_trace()

    
    
    print('Inside sequential_superlong. about to clean doc')
    clean_document(doc)
    # Compile the full report
    


    #doc.save(f'/Users/tomolds/Local/transcription/Stamos/{sym}/Formatted_SuperLong_{today_date}_{sym}.docx')
    
    print('Inside sequential_superlong. about to save doc')
    doc.save(docx_file_path)
    
    full_report = "\n\n".join(report_sections.values())

    full_report = clean_text(full_report)

    all_str = all_str + '\n' + " Full Report: " + '\n' + full_report


    #try:
    #    with open(new_file_path, 'w', encoding='utf-8') as file:
    #        file.write(all_str)
    #except:
    #    print('error writing to file')
    #    pdb.set_trace()

    #print(f"Sequential SuperLong saved to {new_file_path}")

    #pdb.set_trace()

    return  docx_file_path


from fastapi.responses import JSONResponse
@app.post("/api/query/")
async def query_index(query: dict):
    selected_databases = query['databases']
    selected_template = query['template']
    question = query['question']

    db_name = selected_databases[0]
    sym = db_name.split('_')[-1].split('.')[0]

    if selected_template == "SuperLong":
        # Handle SuperLong template separately
        all_chunks = []
        print('in SuperLong section, about to get chunks')
        for db_name in selected_databases:
            sym = db_name.split('_')[-1].split('.')[0]
            chunks, _ = get_chunks_from_db(db_name)  # Only retrieve chunks
            all_chunks.extend(chunks)
        print('in SuperLong section, about to join chunks')
        documents = "\n".join(all_chunks)


        company_name = sym  # Adjust as necessary
        model_name = "gpt-4o"

        print('in SuperLong section, about to call sequential_superlong')
        file_path = sequential_superlong(documents, company_name, model_name)
        print('in SuperLong section, FINISHED sequential_superlong')

        # Return the file path as a response
        return {"file_path": file_path}
        

    
    elif selected_template == "Tear Sheet":
        # Handle SuperLong template separately
        all_chunks = []
        print('in Tear Sheet section, about to get chunks')
        for db_name in selected_databases:
            sym = db_name.split('_')[-1].split('.')[0]
            chunks, _ = get_chunks_from_db(db_name)  # Only retrieve chunks
            all_chunks.extend(chunks)
        print('in Tear Sheet section, about to join chunks')
        documents = "\n".join(all_chunks)


        company_name = sym  # Adjust as necessary
        model_name = "gpt-4o"

        print('in Tear Sheet sectiong section, about to call sequential_tear_sheet_production')
        file_path = sequential_tear_sheet_production(documents, sym, model_name)
        print('in Tear Sheet section section, FINISHED sequential_tear_sheet_production')
        print('file_path:',file_path)

        # Return the file path as a response
        return {"file_path": file_path}


    
    elif selected_template == "Long Form":
        # Handle SuperLong template separately
        all_chunks = []
        print('in Long Form section, about to get chunks')
        for db_name in selected_databases:
            sym = db_name.split('_')[-1].split('.')[0]
            chunks, _ = get_chunks_from_db(db_name)  # Only retrieve chunks
            all_chunks.extend(chunks)
        print('in Long Form section, about to join chunks')
        documents = "\n".join(all_chunks)


        company_name = sym  # Adjust as necessary
        model_name = "gpt-4o"

        print('in Long Form  section, about to call sequential_long_form_production')
        file_path = sequential_long_memo_production(documents, sym, model_name)
        print('in Long Form section, FINISHED sequential_long_form_production. file_path = ',file_path)

        # Return the file path as a response
        return {"file_path": file_path}


    
    elif selected_template == "One Page Current Events":
        # Handle SuperLong template separately
        all_chunks = []
        print('in One Page section, about to get chunks')
        for db_name in selected_databases:
            sym = db_name.split('_')[-1].split('.')[0]
            chunks, _ = get_chunks_from_db(db_name)  # Only retrieve chunks
            all_chunks.extend(chunks)
        print('in One Page section, about to join chunks')
        documents = "\n".join(all_chunks)


        company_name = sym  # Adjust as necessary
        model_name = "gpt-4o"

        print('in One Page section, about to call quick_one_page_production')
        file_path = quick_one_page_production(documents, sym, model_name)
        print('in One Page section, FINISHED quick_one_page_production')

        # Return the file path as a response
        return {"file_path": file_path}



    else:
        # Non-SuperLong case
        index.reset()
        chunks_storage.clear()
        print("Query received:", query)

        # Load the selected databases into the index
        for db_name in selected_databases:
            chunks, embeddings = get_chunks_from_db(db_name)
            for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                index.add(np.array([embedding]))  # Add the embedding to the FAISS index
                chunks_storage[len(chunks_storage)] = chunk  # Store the chunk text

        question_embedding, _ = get_embeddings(question)
        question_embedding = question_embedding[0]  # Assuming the query is short and fits in one chunk

        D, I = index.search(np.array([question_embedding]), k=5)  # Search top-k nearest neighbors

        # Retrieve the corresponding text chunks
        relevant_chunks = [get_chunk_by_index(idx) for idx in I[0]]

        # Combine the relevant chunks into a single prompt
        combined_text = " ".join(relevant_chunks)

        # Template-based prompt creation
        #if selected_template == "One Page Current Events":
        #    prompt = f"Question: {question}\n\nContext: {combined_text}\n\nPlease answere with a one page memo to the investment team describing events covered in the context provided, as well as any potential conclusions and action items:"

        #elif selected_template == "Long Form":
        #    prompt = f"Question: {question}\n\nContext: {combined_text}\n\nAnswer in a long form analysis:"

        if selected_template == "Ad Hoc Query":
            prompt = f"Question: {question}\n\nContext: {combined_text}\n\nPlease Answer this ad hoc query in a fullsome way, based on whatever context provided seems most pertinent:"

        else:
            return {"error": "Invalid template selected"}

        # Generate the response using GPT-4
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a partner at Goldman Sachs or McKinsey and Company. Please answer questions you are asked using the context supplied."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3500
        )

        answer = response.choices[0].message.content
        #answer = clean_text(answer)

        return {"question": question, "answer": answer}


#@app.get("/")
#async def root():
#    return {"message": "Hello World"}



@app.delete("/api/vector-databases/{database_name}")
async def delete_database(database_name: str):
    try:
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(database_name)

        if blob_client.exists():
            blob_client.delete_blob()
            return {"message": f"Database {database_name} deleted successfully."}
        else:
            raise HTTPException(status_code=404, detail="Database not found.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.get("/api/vector-databases")
#@app.get("/vector-databases")
async def list_vector_databases():
    print("Databases being requested")
    blob_list = blob_service_client.get_container_client(container_name).list_blobs()
    databases = [blob.name for blob in blob_list if blob.name.endswith("_index")]
    print("Databases found:", databases)
    return databases

@app.get("/api/prompt-templates")
async def list_prompt_templates():
    print("Templates being requested")
    with open('/app/backend/prompt_templates.txt', 'r') as f:
        templates = f.read().splitlines()
    return templates


@app.post("/api/upload/")
async def upload_file(file: UploadFile = File(...)):
    
    print("About to upload file:",file.filename)
    file_location = f"uploaded_files/{file.filename}"
    with open(file_location, "wb+") as file_object:
        file_object.write(file.file.read())

    # Upload to Azure Blob Storage
    blob_client = blob_service_client.get_blob_client(container=container_name, blob=file.filename)
    with open(file_location, "rb") as data:
        blob_client.upload_blob(data, overwrite=True)
    
    
    content = ""
    if file.filename.endswith(".txt"):
        with open(file_location, "r") as f:
            content = f.read()
    elif file.filename.endswith(".docx"):
        content = read_docx(file_location)
    elif file.filename.endswith(".pdf"):
        content = extract_text_from_pdf(file_location)
    else:
        raise HTTPException(status_code=400, detail="File format not supported.")

    
    ## Read file content and generate embeddings
    #with open(file_location, "r") as f:
    #    content = f.read()
    embeddings, chunks = get_embeddings(content)
    
    # Add each embedding to the FAISS index and store the chunks
    for idx, embedding in enumerate(embeddings):
        index.add(np.array([embedding]))
        chunks_storage[idx] = chunks[idx]

    index_file = f"{file.filename}_index"
    faiss.write_index(index, index_file)
    blob_client_index = blob_service_client.get_blob_client(container=container_name, blob=index_file)
    with open(index_file, "rb") as data:
        blob_client_index.upload_blob(data, overwrite=True)

    chunks_file = f"{file.filename}_chunks"
    with open(chunks_file, "w") as f:
        f.write("\n".join(chunks))
    blob_client_chunks = blob_service_client.get_blob_client(container=container_name, blob=chunks_file)
    with open(chunks_file, "rb") as data:
        blob_client_chunks.upload_blob(data, overwrite=True)

    # Clean up local files
    os.remove(file_location)
    os.remove(index_file)
    os.remove(chunks_file)

    return {"info": f"file '{file.filename}' saved at '{file_location}' and '{index_file}' also saved. "}




# @app.delete("/vector-databases/{database_name}")
# async def delete_database(database_name: str):
#     database_path = f"path/to/your/databases/{database_name}"
#     try:
#         if os.path.exists(database_path):
#             os.remove(database_path)
#             return {"message": f"Database {database_name} deleted successfully."}
#         else:
#             raise HTTPException(status_code=404, detail="Database not found.")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))



@app.delete("/api/vector-databases/{database_name}")
async def delete_database(database_name: str):
    try:
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(database_name)

        if blob_client.exists():
            blob_client.delete_blob()
            return {"message": f"Database {database_name} deleted successfully."}
        else:
            raise HTTPException(status_code=404, detail="Database not found.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))





@app.get("/api/download")
async def download_file(filename: str):
    try:
        #results_dir = os.getenv('RESULTS_DIR', '/app/backend/results')
        results_dir = os.getenv('RESULTS_DIR', 'results')
        file_path = os.path.join(results_dir, filename)
        
        print("Download function called")  # Logging for debugging
        file_path = file_path.replace('"', '').replace("'", "")
        print(f"Received request to download file: {file_path}")  # Logging for debugging
        if not os.path.exists(file_path):
            print(f"File at path {file_path} does not exist.")  # Logging for debugging
            raise HTTPException(status_code=404, detail="File not found")
        print(f"File at path {file_path} exists. Preparing to send...")  # Logging for debugging
        return FileResponse(file_path, filename=os.path.basename(file_path))
    except Exception as e:
        print(f"Error occurred: {e}")  # Log any exceptions that occur
        raise HTTPException(status_code=500, detail="Internal Server Error")


def quick_one_page_production(documents, sym, model_name):
#sequential_tear_sheet_production(content_to_send, file_path,new_file_path)
    
    today_date = datetime.today().strftime('%Y-%m-%d')

    #sanitized_company_name = re.sub(r'\s+', '_', company_name).replace("'", "").replace('"', "")
    sanitized_company_name = re.sub(r'\s+', '_', sym).replace("'", "").replace('"', "")
    #filename = f"Formatted_SuperLong_{today_date}_{sanitized_company_name}.docx"
    filename = f"Formatted_OnePageCurrent_{today_date}_{sym}.docx"
    print("sanitized_company_name:",sanitized_company_name)
    #filename = f"Formatted_SuperLong_{today_date}_{company_name}.docx"
    #docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"
    docx_file_path = os.path.join(RESULTS_DIR, filename)
    

    
    #pdb.set_trace()
    
    all_str = " "



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey  and Company. in this case, you are helping a user create an
             Please create a one page memo to the investment team describing events covered in the context provided, as well as any potential conclusions and action items. the content is : {documents}"
               """}
        ]
          
            
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})

    for attempt in range(5):  # Try up to 5 times
        try:
            response = client.chat.completions.create(
            messages=jls_extract_var,
            model = model_name,
                    temperature = .3,
                    #model="gpt-4",
                
                    max_tokens=3500
)

            break  # If the request was successful, break the loop
        except Exception as e:
            print(f"Attempt {attempt+1} failed with error: {e}")
            if attempt < 4:  # Wait for 2 seconds unless this was the last attem
                time.sleep(2)
                continue  # Try again
            else:
                print("All attempts failed. Exiting.")
                pdb.set_trace()
                #return None
    
   
    response_str = response.choices[0].message.content
   


    # Splitting the string into lines
    one_page = response_str.split('\n')
    one_page_str = '\n'.join(one_page)


    one_page_str = clean_text(one_page_str)

# Open the file in write mode ('w') and write the response_str to it
    
    doc = Document()
 
    # Generate each section and store the results
    formatted_section = analyze_text_with_gpt(one_page_str)    
    add_formatted_content(doc, formatted_section)
    #text_so_far = text_so_far + '\n' + report_sections[section]
        #pdb.set_trace()

    clean_document(doc)


    print('Inside quick_one_page_production. about to clean doc')
    clean_document(doc)
  
    
    print('Inside quick_one_page_production. about to save doc')
    doc.save(docx_file_path)
    print('Inside quick_one_page_production. saved doc to:',docx_file_path)

    return docx_file_path








def test_sequential_tear_sheet_production(chunk, sym, model_name):
#sequential_tear_sheet_production(content_to_send, file_path,new_file_path)
    
    today_date = datetime.today().strftime('%Y-%m-%d_%H-%M-%S')

    #sanitized_company_name = re.sub(r'\s+', '_', company_name).replace("'", "").replace('"', "")
    sanitized_company_name = re.sub(r'\s+', '_', sym).replace("'", "").replace('"', "")
    #filename = f"Formatted_SuperLong_{today_date}_{sanitized_company_name}.docx"
    filename = f"Formatted_TearSheet_{today_date}_{sym}.docx"
    print("sanitized_company_name:",sanitized_company_name)
    #filename = f"Formatted_SuperLong_{today_date}_{company_name}.docx"
    docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"
    print('docx_file_path:',docx_file_path)

    #return docx_file_path    

# Splitting the filename on underscores and extracting the part before '.txt'
    extracted_string = filename.split('_')[-1].split('.')[0]
    
    
    #sym = parts[1]

    sym = extracted_string
    
    
    #pdb.set_trace()
    
    all_str = " "



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
             The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}". 
            for this iteration of the process you should just look at the 'I. Company Thesis' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the history, development, positioning of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief reference to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
            Please don't use overly hyperbolic language, but the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. 
             Overall,  I would like you to tell me about the company
              and its prospects and things you find admirable about its philosphy, history and approach.
                It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied.
             
              also , here are a few other points to keep in mind: {eti_prompt}.
              Where summaries make include observations or conclusions, please make sure to include at least brief reference to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
              in this particular iteration of the process we only want to produce an overal Company Thesis section for the investment analysis summary. This should be similar in style and content to that privoded in the example investment analysis summary. this should be between 250 and 350 words long  AT MOST.
              the Company Thesis section is wehre we talk a little about the origin, history, development, philosophy and mission fo the company.
               the TEXT IS: {chunk}.
              I prefer not to have too many financial results in the ‘Company Thesis’ section  High level, very salient financials are ok in the Company Thesis section, but it should be in service of painting the big picture about the company, rather than a report of recent financials. The Company Thesis section is meant to be a picture picture view of why we might want to be invested in this company.
              In addition, please try to add a sense of inspriration and aspiration to the Company Thesis section. But definitely don't go overboard with enthusiasm.
              Please only output the Company Thesis section in this reponse.  please try hard not be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
               Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})
        

    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content
   


    # Splitting the string into lines
    thesis = response_str.split('\n')
    thesis_str = '\n'.join(thesis)


    thesis_str = clean_text(thesis_str)

    all_str = all_str + '\n' +  thesis_str # prevent double statement of section title
   
 ########################################################################################################################


    ##############################################################################################################################

    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}". for this iteration of the process you should just look at the 'II. Company Basics' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and most important recent financial results of a company we are considering investing in.
          Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             Overall,  I would like you to tell me some basic important facts about the company.  Such as who is the CEO, what is the CEOs tenure, what is the rough size of the company, etx.
              It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              also , here are a few other points to keep in mind: {eti_prompt}.
              in this particular iteration of the process we only want to produce a Company Basics section for the investment analysis summary.  this should be short. absolute maximum 150 words long.  only facts or statistics that are in the text should be included.  please don't include anything in this Company Basics section that you've arlead included in the Company Thesis section.
              I definitely do want to include a sentence about who is the current CEO and what is their tenure in that role, if you have that information.
              The Company Basics section should be short. Concise accurate information.  Maximum of 150 words.
              the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}
              Please only output the Company Basics section in this reponse.  please try hard very hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
              Again, the seciton you are creaing here is the Company Basics section. Previous sections you  want to be sure not be repetitive with in this section is Company thesis: {thesis_str}. Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content

    # Splitting the string into lines
    recent_finanical_results = response_str.split('\n')
    financial_str = '\n'.join(recent_finanical_results)

    financial_str = clean_text(financial_str)

    #all_str = all_str + '\n' + " Company Basics: " + '\n' + financial_str

   
    all_str = all_str +  '\n' + financial_str # prevent double statement of section title


  


    #########################################################################################################################




    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}".
            
             for this iteration of the process you should just look at the 'III. Investment Highlights' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of teh positioning and recent financial results of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
           Please make sure to capture and cite any information related to activist shareholders or investors being involved.  This is always interesting.  Also please make sure to mention any stock buyback or repurchase activity the company may be involved in.
            Please don't use overly hyperbolic language, but
             definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
             Overall,  I would like you to tell me about the company
              and its prospects and things you like and things you would be concerned about as I think about investing in 
              this company.              
                It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              in this particular instance we only want to produce an Investment Highlights section for the blog post.  this should be between 300 and 400 words long.  
              the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}.please don't include anything in this Investment Highlights section that you've arlead included in the Company Thesis section.
                the Company Basics section of this investment analysis summary that was created during a previous call to the model is: {financial_str}. please also don't include anything in this Investment Highlights section that you've arlead included in the Company Basics section.
                Please only output the Investment Highights section in this particular reponse.  please try hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
              Also please try not to use the word 'delve'.
               """}
        ]
     
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content

    # Splitting the string into lines
    investment_highlights = response_str.split('\n')

    highlights_str = '\n'.join(investment_highlights)

    highlights_str = clean_text(highlights_str)

    #all_str = all_str + '\n'+" Investment Highlights " + '\n'+ highlights_str

    all_str = all_str + '\n'+ highlights_str # prevent double statement of section title


    #doc = Document()
 
    ## Generate each section and store the results
    #formatted_section = analyze_text_with_gpt(all_str)    
    #add_formatted_content(doc, formatted_section)
    ##text_so_far = text_so_far + '\n' + report_sections[section]
    #    #pdb.set_trace()

   


    #print('Inside sequential_tear_sheet_production. about to clean doc')
    #clean_document(doc)
  
    
    #print('Inside sequential_tear_sheet_production. about to save doc')
    #doc.save(docx_file_path)
    #print('Inside sequential_tear_sheet_production. saved doc to:',docx_file_path)
    #return docx_file_path 

     
    #doc = Document()
 
    ## Generate each section and store the results
    #formatted_section = analyze_text_with_gpt(all_str)    
    #add_formatted_content(doc, formatted_section)
    ##text_so_far = text_so_far + '\n' + report_quersections[section]
    #    #pdb.set_trace()

   


    #print('Inside sequential_tear_sheet_production. about to clean doc')
    #clean_document(doc)
  
    
    #print('Inside sequential_tear_sheet_production. about to save doc')
    #doc.save(docx_file_path)
    #print('Inside sequential_tear_sheet_production. saved doc to:',
    #docx_file_path)
    #return docx_file_path 

    
   


    #########################################################################################################################



    #jls_extract_var = [
    #        {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
    #         investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
    #        by the model, we will be processing the text sequentially and iteratively.
           
            
    #        The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
    #        here is the example: {stamos_example}". for this iteration of the process you should just look at the 'IV. Investment Concerns' section of the example.
    #        this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
    #        investment analysis summary should appear in the summary we are creating here. keep in mind investment analysis summary we are building here concerns the company described in the text: {sym}
    #        The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and recent financial results of a company we are considering investing in.
    #        Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
           
    #        Please don't use overly hyperbolic language, but
    #         definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
    #         should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
    #         Overall,  in this section I would like you to tell me about the things you would be concerned about as I think about investing in 
    #          this company.  
    #            It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied.
             
             
             
    #          also , here are a few other points to keep in mind: {eti_prompt}.
              
              
    #          in this particular instance we only want to produce an Investment Concerns section for the investment analysis summary.  this should be between 150 and 300 words long.
    #          please tell me what you think might be some risks to the future prospects of the company and why.  Please make sure to provide specific data or references to ground concerns you may have.  we don't want any generic fluff. thank you.  
    #          in general for transcripts of earnings calls, analysts attening the call wil ask questions after the management of the company makes its presentation.  
    #          analysts are often asking about items they are particularly concerned about or intereted in,  typically analyst questions help highligh the most important concerns and also the biggest opportunities.
    #          the text here includes earnigns call  transcripts. please be sure to pay attention to the questions that are asked by the analysts as they can provide important insights into what the market is thinking about the company.
    #          and of course the answers management gives to the questions are also a good source of insights.
              
             
    #          the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}. please don't include anything in this Investment Concerns section that you've arlead included in the Company Thesis section.
    #            the Company Basics section of this investment analysis summary that was created during a previous call to the model is: {financial_str}. please also don't include anything in this Investment Concerns section that you've arlead included in the Company Basics section.
    #            the Investment Highlights section of this investment analysis summarythat was created during a previous call to the model is: {highlights_str} . please don't include anything in this Investment Concerns section that you've arlead included in the Company Thesis section.
    #            Please only output the Investment Concerns section in this reponse.  please try hard Not to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
    #           For the Investment Concerns section you should list what you see as meaningful areas of risk.  But please do NOT end this section with a summarizing statement about your overall risk assessment.
               
    #           Also please try not to use the word 'delve'.
    #           """}
    #    ]
    
    #jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    #response =client.chat.completions.create(
    #    messages=jls_extract_var,
    #    model = model_name,
    #    temperature = .3,
    #    max_tokens=2500
    #)
    
    #response_str = response.choices[0].message.content

    ## Splitting the string into lines
    #concerns = response_str.split('\n')
    #concerns_str = '\n'.join(concerns)

    #concerns_str = clean_text(concerns_str)

    ##all_str = all_str + '\n'+" Investment Concerns: " + '\n'+ concerns_str

    #all_str = all_str + '\n'+ concerns_str # prevent double statement of section title

    
    
    
    
    
    
    
    #doc = Document()
 
    ## Generate each section and store the results
    #formatted_section = analyze_text_with_gpt(all_str)    
    #add_formatted_content(doc, formatted_section)
    ##text_so_far = text_so_far + '\n' + report_sections[section]
    #    #pdb.set_trace()

    #clean_document(doc)


    #print('Inside sequential_tear_sheet_production. about to clean doc')
    #clean_document(doc)
  
    
    #print('Inside sequential_tear_sheet_production. about to save doc')
    #doc.save(docx_file_path)
    #print('Inside sequential_tear_sheet_production. saved doc to:',
    #docx_file_path)
    #return docx_file_path 









def sequential_tear_sheet_production(chunk, sym, model_name):
#sequential_tear_sheet_production(content_to_send, file_path,new_file_path)
    
    today_date = datetime.today().strftime('%Y-%m-%d')

    #sanitized_company_name = re.sub(r'\s+', '_', company_name).replace("'", "").replace('"', "")
    sanitized_company_name = re.sub(r'\s+', '_', sym).replace("'", "").replace('"', "")
    #filename = f"Formatted_SuperLong_{today_date}_{sanitized_company_name}.docx"
    filename = f"Formatted_TearSheet_{today_date}_{sym}.docx"
    print("sanitized_company_name:",sanitized_company_name)
    #filename = f"Formatted_SuperLong_{today_date}_{company_name}.docx"
    #docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"
    docx_file_path = os.path.join(RESULTS_DIR, filename)
    print('docx_file_path:',docx_file_path)

    #return docx_file_path    

# Splitting the filename on underscores and extracting the part before '.txt'
    extracted_string = filename.split('_')[-1].split('.')[0]
    
    
    #sym = parts[1]

    sym = extracted_string
    
    
    #pdb.set_trace()
    
    all_str = " "



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
             The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}". 
            for this iteration of the process you should just look at the 'I. Company Thesis' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the history, development, positioning of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief reference to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
            Please don't use overly hyperbolic language, but the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. 
             Overall,  I would like you to tell me about the company
              and its prospects and things you find admirable about its philosphy, history and approach.
                It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied.
             
              also , here are a few other points to keep in mind: {eti_prompt}.
              Where summaries make include observations or conclusions, please make sure to include at least brief reference to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
              in this particular iteration of the process we only want to produce an overal Company Thesis section for the investment analysis summary. This should be similar in style and content to that privoded in the example investment analysis summary. this should be between 250 and 350 words long  AT MOST.
              the Company Thesis section is wehre we talk a little about the origin, history, development, philosophy and mission fo the company.
               the TEXT IS: {chunk}.
              I prefer not to have too many financial results in the ‘Company Thesis’ section  High level, very salient financials are ok in the Company Thesis section, but it should be in service of painting the big picture about the company, rather than a report of recent financials. The Company Thesis section is meant to be a picture picture view of why we might want to be invested in this company.
              In addition, please try to add a sense of inspriration and aspiration to the Company Thesis section. But definitely don't go overboard with enthusiasm.
              Please only output the Company Thesis section in this reponse.  please try hard not be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
               Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})
        

    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content
   


    # Splitting the string into lines
    thesis = response_str.split('\n')
    thesis_str = '\n'.join(thesis)


    thesis_str = clean_text(thesis_str)

    all_str = all_str + '\n' +  thesis_str # prevent double statement of section title
   

    ##############################################################################################################################

    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}". for this iteration of the process you should just look at the 'II. Company Basics' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and most important recent financial results of a company we are considering investing in.
          Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             Overall,  I would like you to tell me some basic important facts about the company.  Such as who is the CEO, what is the CEOs tenure, what is the rough size of the company, etx.
              It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              also , here are a few other points to keep in mind: {eti_prompt}.
              in this particular iteration of the process we only want to produce a Company Basics section for the investment analysis summary.  this should be short. absolute maximum 150 words long.  only facts or statistics that are in the text should be included.  please don't include anything in this Company Basics section that you've arlead included in the Company Thesis section.
              I definitely do want to include a sentence about who is the current CEO and what is their tenure in that role, if you have that information.
              The Company Basics section should be short. Concise accurate information.  Maximum of 150 words.
              the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}
              Please only output the Company Basics section in this reponse.  please try hard very hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
              Again, the seciton you are creaing here is the Company Basics section. Previous sections you  want to be sure not be repetitive with in this section is Company thesis: {thesis_str}. Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content

    # Splitting the string into lines
    recent_finanical_results = response_str.split('\n')
    financial_str = '\n'.join(recent_finanical_results)

    financial_str = clean_text(financial_str)

    #all_str = all_str + '\n' + " Company Basics: " + '\n' + financial_str

   
    all_str = all_str +  '\n' + financial_str # prevent double statement of section title

    ########################################################################################################################



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}".
            
             for this iteration of the process you should just look at the 'III. Investment Highlights' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of teh positioning and recent financial results of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
           Please make sure to capture and cite any information related to activist shareholders or investors being involved.  This is always interesting.  Also please make sure to mention any stock buyback or repurchase activity the company may be involved in.
            Please don't use overly hyperbolic language, but
             definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
             Overall,  I would like you to tell me about the company
              and its prospects and things you like and things you would be concerned about as I think about investing in 
              this company.              
                It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              in this particular instance we only want to produce an Investment Highlights section for the blog post.  this should be between 300 and 400 words long.  
              the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}.please don't include anything in this Investment Highlights section that you've arlead included in the Company Thesis section.
                the Company Basics section of this investment analysis summary that was created during a previous call to the model is: {financial_str}. please also don't include anything in this Investment Highlights section that you've arlead included in the Company Basics section.
                Please only output the Investment Highights section in this particular reponse.  please try hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
              Also please try not to use the word 'delve'.
               """}
        ]
     
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content

    # Splitting the string into lines
    investment_highlights = response_str.split('\n')

    highlights_str = '\n'.join(investment_highlights)

    highlights_str = clean_text(highlights_str)

    #all_str = all_str + '\n'+" Investment Highlights " + '\n'+ highlights_str

    all_str = all_str + '\n'+ highlights_str # prevent double statement of section title


    ########################################################################################################################



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            
            The user will provide you with an example of a recent investment analysis summary so you can see the kind of information we are going after. 
            here is the example: {stamos_example}". for this iteration of the process you should just look at the 'IV. Investment Concerns' section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. keep in mind investment analysis summary we are building here concerns the company described in the text: {sym}
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and recent financial results of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
           
            Please don't use overly hyperbolic language, but
             definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
             Overall,  in this section I would like you to tell me about the things you would be concerned about as I think about investing in 
              this company.  
                It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied.
             
             
             
              also , here are a few other points to keep in mind: {eti_prompt}.
              
              
              in this particular instance we only want to produce an Investment Concerns section for the investment analysis summary.  this should be between 150 and 300 words long.
              please tell me what you think might be some risks to the future prospects of the company and why.  Please make sure to provide specific data or references to ground concerns you may have.  we don't want any generic fluff. thank you.  
              in general for transcripts of earnings calls, analysts attening the call wil ask questions after the management of the company makes its presentation.  
              analysts are often asking about items they are particularly concerned about or intereted in,  typically analyst questions help highligh the most important concerns and also the biggest opportunities.
              the text here includes earnigns call  transcripts. please be sure to pay attention to the questions that are asked by the analysts as they can provide important insights into what the market is thinking about the company.
              and of course the answers management gives to the questions are also a good source of insights.
              
             
              the TEXT IS: {chunk} and the Company Thesis section of this investment analysis summary that was created during a previous call to the model is: {thesis_str}. please don't include anything in this Investment Concerns section that you've arlead included in the Company Thesis section.
                the Company Basics section of this investment analysis summary that was created during a previous call to the model is: {financial_str}. please also don't include anything in this Investment Concerns section that you've arlead included in the Company Basics section.
                the Investment Highlights section of this investment analysis summarythat was created during a previous call to the model is: {highlights_str} . please don't include anything in this Investment Concerns section that you've arlead included in the Company Thesis section.
                Please only output the Investment Concerns section in this reponse.  please try hard Not to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
               For the Investment Concerns section you should list what you see as meaningful areas of risk.  But please do NOT end this section with a summarizing statement about your overall risk assessment.
               
               Also please try not to use the word 'delve'.
               """}
        ]
    
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content

    # Splitting the string into lines
    concerns = response_str.split('\n')
    concerns_str = '\n'.join(concerns)

    concerns_str = clean_text(concerns_str)

    #all_str = all_str + '\n'+" Investment Concerns: " + '\n'+ concerns_str

    all_str = all_str + '\n'+ concerns_str # prevent double statement of section title

 

# Open the file in write mode ('w') and write the response_str to it
    
    doc = Document()
 
    # Generate each section and store the results
    formatted_section = analyze_text_with_gpt(all_str)    
    add_formatted_content(doc, formatted_section)
    #text_so_far = text_so_far + '\n' + report_sections[section]
        #pdb.set_trace()

    clean_document(doc)


    print('Inside sequential_tear_sheet_production. about to clean doc')
    clean_document(doc)
  
    
    print('Inside sequential_tear_sheet_production. about to save doc')
    doc.save(docx_file_path)
    print('Inside sequential_tear_sheet_production. saved doc to:',
    docx_file_path)
    return  docx_file_path





def test_sequential_long_memo_production(chunk, sym, model_name):
#sequential_tear_sheet_production(content_to_send, file_path,new_file_path)
     
    
    print('Inside test_sequential_long_memo_production')
    today_date = datetime.today().strftime('%Y-%m-%d_%H-%M-%S')

    print('today_date:',today_date)

   
    sanitized_company_name = re.sub(r'\s+', '_', sym).replace("'", "").replace('"', "")
  
    filename = f"Formatted_LongForm_{today_date}_{sym}.docx"
    print("sanitized_company_name:",sanitized_company_name)
  
    docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"

    #return docx_file_path

    
    all_str = " "



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
             The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
            here are the two examples: {long_form_examples}". 
            for this iteration of the process you should just look at the Company Origin and the Why (Mission, Vision, and Values )section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the history, development, positioning of a company we are considering investing in.
             Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             And things you find admirable about its philosphy, histroy and approach. Again, the Why sections hould contain your best efforts at summarizing this Mission, Vision and Values of the company.
            
            Please don't use overly hyperbolic language, but the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. 
             Overall,  I would like you to tell me about the company
              and its origin and history.
                It is very important for purposes of these analyses that you only use and report facts, statistics and numbers which exist in the text that is supplied.
             
              also , here are a few other points to keep in mind: {eti_prompt}.
              in this particular iteration of the process we only want to produce an overal Company Origin section for the investment analysis summary. This should be similar in style and content to that privoded in the examples investment analysis summary. this should be no longer than 300 or 400 words long. 
              the Company Origin section is wehre we talk a little about the origin, history, development, philosophy and mission fo the company.
              I definitely want to include an explicit  "Why" component in this section.  The "Why" component should be a brief summary of the company's Mission, Vision, and Values to the best of your ability, or similar.  The "Why" component should be no longer than 100 words. 
               the TEXT IS: {chunk}.
              In addition, please try to add a sense of inspriration and aspiration to the Company Origin section. But definitely don't go overboard with enthusiasm.
              Please only output the Company Thesis section in this reponse.  please try hard not be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
               Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
               Also please try not to use the word 'delve'.
               """}
        ]
    
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})
 

    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content


    # Splitting the string into lines
    origin = response_str.split('\n')
    origin_str = '\n'.join(origin)

    origin_str = clean_text(origin_str)


    all_str = all_str + '\n'  + origin_str # prevent double statement of section title



    doc = Document()

    formatted_section = analyze_text_with_gpt(all_str)    
    add_formatted_content(doc, formatted_section)

    print('Inside test_sequential_long_memo_production. about to clean doc')
    clean_document(doc)
  
    
    print('Inside test_sequential_long_memo_production. about to save doc')
    doc.save(docx_file_path)
    print('Inside test_sequential_long_memo_production. saved doc to:',docx_file_path)
    return docx_file_path 
 



    ##############################################################################################################################

    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
            here is the example: {long_form_examples}". for this iteration of the process you should just look at the 'What' section of the examples.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and most important recent financial results of a company we are considering investing in.
          Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             Overall,  I would like you to tell me some basic important facts about the company.  This mainly relates to financial results, line of busines focused on , geographies covered , etc.
               It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              also , here are a few other points to keep in mind: {eti_prompt}.
              
              Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
              in this particular iteration of the process we only want to produce a 'What' section for the investment analysis summary.  this should be between roughly 400 to 600 words long.  
              
              the TEXT IS: {chunk} and the Company Origin section of this investment analysis summary that was created during a previous call to the model is: {origin_str}
              Please only output the 'What' section in this reponse.  please try hard very hard NOT to be repetitive either within or across the sections you are writing.  
              Again, the seciton you are creaing here is the 'What' section. Previous sections you  want to be sure not be repetitive with in this section is Company Origin: {origin_str}. Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlines in the system content above."})
 
    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content


    # Splitting the string into lines
    what_section = response_str.split('\n')
    what_str = '\n'.join(what_section)

    what_str = clean_text(what_str)

   
    all_str = all_str + '\n' +  what_str   


    
    doc = Document()

    formatted_section = analyze_text_with_gpt(all_str)    
    add_formatted_content(doc, formatted_section)

    print('Inside test_sequential_long_memo_production. about to clean doc')
    clean_document(doc)
  
    
    print('Inside test_sequential_long_memo_production. about to save doc')
    doc.save(docx_file_path)
    print('Inside test_sequential_long_memo_production. saved doc to:',
    docx_file_path)
    return docx_file_path 
 

    ########################################################################################################################


    #jls_extract_var = [
    #        {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
    #         investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
    #        by the model, we will be processing the text sequentially and iteratively.
           
    #        The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
    #        here is the example: {long_form_examples}".
    #        these examples are only to show you the format and style of what we are trying to produce - none of the content in this example
    #        investment analysis summary should appear in the summary we are creating here.
    #         for this iteration of the process you should just look at the 'How' section of the example.
    #         this largely relates to organization, information about senior executires , leadership and coporate governance. also informaiton about how the company's operations are structured if possible. and how its cusomters and suppliers are segmented and managed, if possible.
            
    #         this investment analysis summary concerns the company described in the text: {sym}.
    #        The overall purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and recent financial results of a company we are considering investing in.
    #        Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
    #        Please make sure to capture and cite any information related to activist shareholders or investors being involved.  This is always interesting.  Also please make sure to mention any stock buyback or repurchase activity the company may be involved in.
    #        Please don't use overly hyperbolic language, but
    #         definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
    #         should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
    #         Overall,  I would like you to tell me about the company
    #          and its prospects and things you like and things you would be concerned about as I think about investing in 
    #          this company.              
    #             It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
    #          in this particular instance we only want to produce an Investment Highlights section for the blog post.  this should be between 300 and 400 words long.  
    #          the TEXT IS: {chunk} and the Company origin  section of this investment analysis summary that was created during a previous call to the model is: {origin_str}.please don't include anything in this Investment Highlights section that you've arlead included in the Company Thesis section.
    #            the 'What' section of this investment analysis summary that was created during a previous call to the model is: {what_str}. please also don't include anything in this Investment Highlights section that you've arlead included in the Company Basics section.
    #            Please only output the 'How' section in this particular reponse.  please try hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
    #          Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
    #          Also please try not to use the word 'delve'.
    #           """}
    #    ]
    #jls_extract_var.append({"role": "user", "content": f"please execute the task outlines in the system content above."})


    #response =client.chat.completions.create(
    #    messages=jls_extract_var,
    #    model = model_name,
    #    temperature = .3,
    #    max_tokens=2500
    #)
    
    #response_str = response.choices[0].message.content


    ## Splitting the string into lines
    #how_section = response_str.split('\n')

    #how_str = '\n'.join(how_section)

    #how_str = clean_text(how_str)

    ##all_str = all_str + '\n'+" How " + '\n'+ how_str

    #all_str = all_str + '\n'+ how_str


    #doc = Document()

    #print('Inside test_sequential_long_memo_production. about to clean doc')
    #clean_document(doc)
  
    
    #print('Inside test_sequential_long_memo_production. about to save doc')
    #doc.save(docx_file_path)
    #print('Inside test_sequential_long_memo_production. saved doc to:',
    #docx_file_path)
    #return docx_file_path 






def sequential_long_memo_production(chunk, sym, model_name):
#sequential_tear_sheet_production(content_to_send, file_path,new_file_path)
     
    today_date = datetime.today().strftime('%Y-%m-%d')

   
    sanitized_company_name = re.sub(r'\s+', '_', sym).replace("'", "").replace('"', "")
  
    filename = f"Formatted_LongForm_{today_date}_{sym}.docx"
    print("sanitized_company_name:",sanitized_company_name)
  

    #docx_file_path = f"/home/azureuser/charlie/backend/results/{filename}"
    docx_file_path = os.path.join(RESULTS_DIR, filename)
    #return docx_file_path

    
    all_str = " "



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
             The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
            here are the two examples: {long_form_examples}". 
            for this iteration of the process you should just look at the Company Origin and the Why (Mission, Vision, and Values )section of the example.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the history, development, positioning of a company we are considering investing in.
             Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             And things you find admirable about its philosphy, histroy and approach. Again, the Why sections hould contain your best efforts at summarizing this Mission, Vision and Values of the company.
            
            Please don't use overly hyperbolic language, but the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. 
             Overall,  I would like you to tell me about the company
              and its origin and history.
                It is very important for purposes of these analyses that you only use and report facts, statistics and numbers which exist in the text that is supplied.
             
              also , here are a few other points to keep in mind: {eti_prompt}.
              in this particular iteration of the process we only want to produce an overal Company Origin section for the investment analysis summary. This should be similar in style and content to that privoded in the examples investment analysis summary. this should be no longer than 300 or 400 words long. 
              the Company Origin section is wehre we talk a little about the origin, history, development, philosophy and mission fo the company.
              I definitely want to include an explicit  "Why" component in this section.  The "Why" component should be a brief summary of the company's Mission, Vision, and Values to the best of your ability, or similar.  The "Why" component should be no longer than 100 words. 
               the TEXT IS: {chunk}.
              In addition, please try to add a sense of inspriration and aspiration to the Company Origin section. But definitely don't go overboard with enthusiasm.
              Please only output the Company Thesis section in this reponse.  please try hard not be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
               Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
               Also please try not to use the word 'delve'.
               """}
        ]
    
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlined in the system content above."})
 

    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content


    # Splitting the string into lines
    origin = response_str.split('\n')
    origin_str = '\n'.join(origin)

    origin_str = clean_text(origin_str)


    all_str = all_str + '\n'  + origin_str # prevent double statement of section title


    ##############################################################################################################################

    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
            here is the example: {long_form_examples}". for this iteration of the process you should just look at the 'What' section of the examples.
            this example investment analysis summary is only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here. this investment analysis summary concerns the company described in the text: {sym}.
            The purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and most important recent financial results of a company we are considering investing in.
          Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
             Overall,  I would like you to tell me some basic important facts about the company.  This mainly relates to financial results, line of busines focused on , geographies covered , etc.
               It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              also , here are a few other points to keep in mind: {eti_prompt}.
              
              Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
              in this particular iteration of the process we only want to produce a 'What' section for the investment analysis summary.  this should be between roughly 400 to 600 words long.  
              
              the TEXT IS: {chunk} and the Company Origin section of this investment analysis summary that was created during a previous call to the model is: {origin_str}
              Please only output the 'What' section in this reponse.  please try hard very hard NOT to be repetitive either within or across the sections you are writing.  
              Again, the seciton you are creaing here is the 'What' section. Previous sections you  want to be sure not be repetitive with in this section is Company Origin: {origin_str}. Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlines in the system content above."})
 
    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content


    # Splitting the string into lines
    what_section = response_str.split('\n')
    what_str = '\n'.join(what_section)

    what_str = clean_text(what_str)

   
    all_str = all_str + '\n' +  what_str    

    ########################################################################################################################



    jls_extract_var = [
            {"role": "system", "content": f"""You are a helpful assistant. You have the business acumen and drive to discover value creation opportunities of a 15 year partner at Goldman Sachs or McKinsey and Company. in this case, you are helping a user create an
             investment analysis summary based on a series of recent SEC filings and some broader internet research. due to limitations in the length of output tokens that can be produced 
            by the model, we will be processing the text sequentially and iteratively.
           
            The user will provide you with two examples of recent investment analysis summaries so you can see the kind of information we are going after. 
            here is the example: {long_form_examples}".
            these examples are only to show you the format and style of what we are trying to produce - none of the content in this example
            investment analysis summary should appear in the summary we are creating here.
             for this iteration of the process you should just look at the 'How' section of the example.
             this largely relates to organization, information about senior executires , leadership and coporate governance. also informaiton about how the company's operations are structured if possible. and how its cusomters and suppliers are segmented and managed, if possible.
            
             this investment analysis summary concerns the company described in the text: {sym}.
            The overall purpose of the investment analysis summary is to provide portfolio managers a summary of the positioning and recent financial results of a company we are considering investing in.
            Where summaries make include observations or conclusions, please make sure to include at least brief references to specific data which leads to these conclusions of observations.  we never want to be accused of speaking in generalities which might equally apply to any company in an industry .  supporting specifics are always desired. 
            Please make sure to capture and cite any information related to activist shareholders or investors being involved.  This is always interesting.  Also please make sure to mention any stock buyback or repurchase activity the company may be involved in.
            Please don't use overly hyperbolic language, but
             definitely do recognize and mention financial successes, capturing and citing data and statistics where possible. the tone 
             should be more cautiously optimistic, where optimisim is appropriate, as befits analyses related to potential investments. if negativity is appropriate, be negative, please be sure to back it up with data and analysis.  
             Overall,  I would like you to tell me about the company
              and its prospects and things you like and things you would be concerned about as I think about investing in 
              this company.              
                 It is very important for purposes of these analyses that you only use and report statistics and numbers which exist in the text that is supplied. Don't make anyhthing up. 
              in this particular instance we only want to produce an Investment Highlights section for the blog post.  this should be between 300 and 400 words long.  
              the TEXT IS: {chunk} and the Company origin  section of this investment analysis summary that was created during a previous call to the model is: {origin_str}.please don't include anything in this Investment Highlights section that you've arlead included in the Company Thesis section.
                the 'What' section of this investment analysis summary that was created during a previous call to the model is: {what_str}. please also don't include anything in this Investment Highlights section that you've arlead included in the Company Basics section.
                Please only output the 'How' section in this particular reponse.  please try hard NOT to be repetitive either within or across the sections you are writing.  if this means that some sections are shorter than the target length, that is fine.
              Please do not include a summarizing or concluding paragraph in this section of the exercise. no ending paragrpah that summarizes overall takeways or inferences from the seciton, please.
              Also please try not to use the word 'delve'.
               """}
        ]
    jls_extract_var.append({"role": "user", "content": f"please execute the task outlines in the system content above."})


    response =client.chat.completions.create(
        messages=jls_extract_var,
        model = model_name,
        temperature = .3,
        max_tokens=2500
    )
    
    response_str = response.choices[0].message.content


    # Splitting the string into lines
    how_section = response_str.split('\n')

    how_str = '\n'.join(how_section)

    how_str = clean_text(how_str)

    #all_str = all_str + '\n'+" How " + '\n'+ how_str

    all_str = all_str + '\n'+ how_str

    
    doc = Document()
 
    # Generate each section and store the results
    formatted_section = analyze_text_with_gpt(all_str)    
    add_formatted_content(doc, formatted_section)
   

    print('Inside sequential_long_memo_production. about to clean doc')
    clean_document(doc)
  
    
    print('Inside sequential_long_memo_production. about to save doc')
    doc.save(docx_file_path)
    print('docx_file_path:',docx_file_path)
   
    return docx_file_path






if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
